from docx import Document

doc = Document('.trae/《主播模拟器：双面人生》剧情策划方案(1).docx')

print("=" * 80)
print("文档内容")
print("=" * 80)

for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"{para.text}")

print("\n" + "=" * 80)
print("表格内容")
print("=" * 80)

for table_idx, table in enumerate(doc.tables):
    print(f"\n--- 表格 {table_idx + 1} ---")
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        print(" | ".join(row_data))
